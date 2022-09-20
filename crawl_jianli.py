#!/usr/bin/env python
# -*- coding: utf-8 -*-

__mtime__ = '2019/10/8'

import os
import random
import re

import requests
from pyquery import PyQuery as pq
from docx import Document

import config as cfg
import utils
import common as com


jianli_filter_word=[
    '千图','网','58pic','com','千小图','qiantu','库宝','千库','588ku',
    'qianku','乔布','乔小布','千图网','千库网','小简','极简','风云','风小云',
    'qiantuwang','qiankuwang','58','千 图 网','千 图','千 小 图','千',
    'qiaobutang','taobao','OfficePLUS','包小图','包图网','ibaotu','ibaotu.com',
    'docer','Docer','小千千','Qiantuwang','唐小糖','糖小糖','陈思涵','飞小扬',
    '小熊猫','熊猫','熊小猫','tukuppt','xiongmao','xiaoxiongmao','飞扬','Jianqiqi','jianqiqi',
    'XIAOXIONGMAO','小 熊猫','熊猫办公','沿途旅客','Xiaoxiongmao','pic','Tukuppt','panda','xiong',
    'mao','亮亮图文','亮亮','梦创','上官凌','cabablist','lingxian','小 熊 猫','亮亮图文','千团','damod',
]

jianli_filter_word=sorted(jianli_filter_word,key=lambda x:len(x),reverse=True)



def replace_text_jianli(path):
    for root,dirs,files in os.walk(path):
        for f in files:
            try:
                if f.endswith(('docx','doc')):
                    full_path=root+'/'+f
                    print(full_path)
                    doc=Document(full_path)

                    # 文本框过滤

                    children = doc.element.body.iter()
                    tags = []
                    for child in children:
                        # 通过类型判断目录
                        if child.tag.endswith(('AlternateContent','textbox')):
                            for ci in child.iter():
                                tags.append(ci.tag)
                                if ci.tag.endswith(('main}r', 'main}pPr')):
                                    if ci.text!=None:
                                        # print(ci.text)
                                        for filter_word in jianli_filter_word:
                                            if filter_word in ci.text:
                                                replace_to=com.choice_text_to_replace(filter_word)
                                                ci.text=ci.text.replace(filter_word,replace_to)
                                                print('{}:{}过滤完毕'.format(f,'文本框内容'))
                    # 表格里的数据过滤
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    for r in p.runs:
                                        # print(r.text)
                                        for filter_word in jianli_filter_word:
                                            if filter_word in r.text:
                                                replace_to=com.choice_text_to_replace(filter_word)
                                                r.text=r.text.replace(filter_word,replace_to)
                                                print('{}:{}过滤完毕'.format(f,'表格数据'))
                    # # 原始文本过滤
                    for p in doc.paragraphs:
                        for o_r in p.runs:
                            # print(o_r.text)
                            for filter_word in jianli_filter_word:
                                if filter_word in o_r.text:
                                    replace_to=com.choice_text_to_replace(filter_word)
                                    o_r.text=o_r.text.replace(filter_word,replace_to)
                                    print('{}:{}过滤完毕'.format(f,'原始文本'))

                    doc.save(full_path)
            except Exception as e:
                print(e)


def replace_text_origin(doc):
    for p in doc.paragraphs:
        for r in p.runs:
            print(r.text)
            com.replace_text(r.text,jianli_filter_word)


def replace_text_table(doc):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        print(r.text)
                        com.replace_text(r.text,jianli_filter_word)


def replace_text_window(doc):

    children = doc.element.body.iter()
    tags = []
    for child in children:
        # 通过类型判断目录
        if child.tag.endswith(('AlternateContent','textbox')):
            for ci in child.iter():
                tags.append(ci.tag)
                if ci.tag.endswith(('main}r', 'main}pPr')):
                    if ci.text!=None:
                        print(ci.text)
                        com.replace_text(ci.text,jianli_filter_word)



def rename_file(path=com.path):
    li=[' ','空白','下载','免费','的自我介绍']
    for root,dirs,files in os.walk(path):
        for f in files:
            full_path=root+'/'+f
            for i in li:
                if i in f:
                    f=f.replace(i,'')

            file_name,ext=os.path.splitext(f)
            if '简历' not in f:
                f=file_name+'个人简历'+ext

            new_path=root+'/'+f

            res=re.findall('^(\d+).+\.docx',f)
            if res!=[]:
                res=res[0]
                f=f.replace(res,'')
                new_path=root+'/'+f

            if os.path.exists(new_path):

                new_path=root+'/'+file_name+str(random.randint(1,1000))+ext

            os.rename(full_path,new_path)
            print('{}:重名名成功！'.format(new_path))

def rename_filter_num(path=cfg.save_path):
     for root,dirs,files in os.walk(path):
         for f in files:
            full_path=root+'/'+f
            res=re.findall('(\d+)\.docx',f)
            if res!=[]:
                new_name=f.replace(res[0],'')
                new_path=root+'/'+new_name
                os.rename(full_path,new_path)
                print('{}:去除数字完成！'.format(new_name))


def handle_jianli_name(path=com.path):
    for root,dirs,files in os.walk(path):
        for f in files:
            if f.endswith(('docx','.doc')):
                full_path=root+'/'+f
                file_name,ext=os.path.splitext(f)
                if ('（' in file_name or '(' in file_name) and ('）' not in file_name and ')' not in file_name):
                    file_name+='）'
                    new_path=root+'/'+file_name+ext
                    os.rename(full_path,new_path)
                    print('{}:改名成功！'.format(new_path))



if __name__=='__main__':



    # rename_file()

    # handle_jianli_name(com.path)
    # handle_jianli_name(cfg.save_path)
    replace_text_jianli(com.path)
    # replace_text_jianli(cfg.save_path)
    # rename_filter_num()
    # utils.many_doc_to_pdf(cfg.save_path)
    # utils.many_doc_to_pdf(com.path)
    # utils.many_pdf_to_png(cfg.save_path,cfg.pic_path)




