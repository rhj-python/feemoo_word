#!/usr/bin/env python
# -*- coding: utf-8 -*-

__mtime__ = '2019/10/12'

import os

import requests
from pyquery import PyQuery as pq
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

import config as cfg
import utils

base_url='https://wenku.baidu.com'

menu_url='{}/list/195?od=2&pn=0'.format(base_url)



def get_url_li(url):
    li=[]
    res=requests.get(url,headers=cfg.get_headers()).content.decode('gbk')
    html=pq(res)
    datas=html('table#docList>tbody>tr:not(first-of-type)')
    for i in datas.items():
        doc_type=i('td:first-of-type>span').attr('title')
        title=i('td:nth-of-type(2)>a').text()
        final_url=i('td:nth-of-type(2)>a').attr('href')
        if doc_type in ('doc','docx'):
            final_url=base_url+final_url
            # print(doc_type)
            print(title)
            # print(final_url)
            li.append(final_url)
    return li

def save_file(path,li):
    with open(path,'a') as f:
        for i in li:
            f.write(i)
            f.write('\n')
    print('写入完毕')


def re_text_docx(path):
    for root,dirs,files in os.walk(path):
        for f in files:
            if f.endswith(('docx')):
                full_path=root+'/'+f
                doc=Document(full_path)

                p1=doc.paragraphs[0]
                pf=p1.paragraph_format
                pf.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
                run=p1.runs[0]
                run.bold=True
                run.font.size=18*12800

                for p in doc.paragraphs[1:]:
                    print('P:{}'.format(p.text))
                    p.paragraph_format.line_spacing=1.5
                    # for run in p.runs:
                    #     print('run:{}'.format(run.text))
                doc.save(full_path)

def main(start,end):
    for i in range(start,end):

        # 行政规范最高最新发布排序
        # menu_url='{}/list/195?od=1&pn={}'.format(base_url,i*25)


        # 合同最新发布
        # menu_url='{}/list/75?od=1&pn={}'.format(base_url,i*25)

        # 商业计划最新发布
        menu_url='{}/list/182?od=1&pn={}'.format(base_url,i*25)


        li=get_url_li(menu_url)
        save_file('1.txt',li)



if __name__=='__main__':
    # main(0,20)

    path='g:/资源/WORD资源/2019/10月/待处理'

    utils.rename_file(path,cfg.word_filter)
    # utils.rename_file_front(path)
    # utils.rename_file_end(path)

    # utils.rename_file(cfg.save_path,cfg.word_filter)
    # utils.doc_to_docx(cfg.save_path,del_origin_file=True)
    # utils.rename_file_front(cfg.save_path)

    # utils.rename_file_end(cfg.save_path)

    # utils.replace_many(cfg.save_path,cfg.name_filter)
    # re_text_docx(cfg.save_path)

    utils.many_doc_to_pdf(cfg.save_path)
    utils.many_pdf_to_png(cfg.save_path,cfg.pic_path)
