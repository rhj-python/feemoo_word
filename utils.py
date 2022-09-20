#!/usr/bin/env python
# -*- coding: utf-8 -*-

__mtime__ = '2019/10/8'


from io import StringIO,open
import os
import threading
import re
import random
import zipfile

import win32gui
import win32con
import fitz
from win32com import client as wc
from docx import Document
from pathlib import Path

import config as cfg
import common as com


def unzip(file):

    file_name, ext = os.path.splitext(file)
    try:
        if ext == ".zip":
            print ('unzip', file)
        f = zipfile.ZipFile(file)
        # f.extractall(path=file_name.encode(file_encoding)) # 通过path指定解压的路径
        f.extractall(path=r'{}'.format(file_name)) # 通过path指定解压的路径
    except Exception as e:
        print(e)

# 将zip文件解压到其所在目录
def unzip_cn(file,dst=com.path):

    file_name, ext = os.path.splitext(file)
    is_zip = zipfile.is_zipfile(file)

    if is_zip:
        zip_file_contents = zipfile.ZipFile(file, 'r')
        for file in zip_file_contents.namelist():
            filename = file.encode('cp437').decode('gbk')#先使用cp437编码，然后再使用gbk解码

            zip_file_contents.extract(file,dst)#解压缩ZIP文件
            os.chdir(dst)#切换到目标目录
            if not os.path.exists(dst+'/'+filename):
                os.rename(file,filename)#重命名文件
            print ('unzip', file)


# 递归删除目录及目录下文件
def removeDir(dirPath):
    if not os.path.isdir(dirPath):
       return
    files = os.listdir(dirPath)
    try:
       for file in files:
           filePath = os.path.join(dirPath,file)
           if os.path.isfile(filePath):
               os.remove(filePath)
           elif os.path.isdir(filePath):
                removeDir(filePath)
       os.rmdir(dirPath)
    except Exception as e:
       print(e)

def get_changed(path,choice='files',need_ext=True):
    li=[]

    for root,dirs,files in os.walk(path):
        if choice=='files':
            if need_ext==False:
                for f in files:
                    file_name,ext=os.path.splitext(f)
                    li.append(file_name)
            else:
                li.extend(files)
        elif '.' in choice or isinstance(choice,tuple):
                for f in files:
                    if f.endswith(choice):
                        li.append(f)

        elif choice=='dirs':
            li.extend(dirs)

    return li


def win_handle(path,title=u'打开'):
     # win32gui

    dialog = win32gui.FindWindow('#32770', title)  # 对话框
    ComboBoxEx32 = win32gui.FindWindowEx(dialog, 0, 'ComboBoxEx32', None)
    ComboBox = win32gui.FindWindowEx(ComboBoxEx32, 0, 'ComboBox', None)
    Edit = win32gui.FindWindowEx(ComboBox, 0, 'Edit', None)  # 上面三句依次寻找对象，直到找到输入框Edit对象的句柄
    button = win32gui.FindWindowEx(dialog, 0, 'Button', None)  # 确定按钮Button

    # time.sleep(0.5)

    win32gui.SendMessage(Edit, win32con.WM_SETTEXT, None,path)  # 往输入框输入绝对地址
    win32gui.SendMessage(dialog, win32con.WM_COMMAND, 1, button)  # 按button


def get_file_path(dir_name,end_with=(".docx",'.doc'),to_win32=False):
    li=[]
    for root,dirs,files in os.walk(dir_name):
        for file_name in files:
            if file_name.endswith(end_with):
                full_path=os.path.join(root,file_name)
                li.append(full_path)
    if to_win32==True:
        li2=[]
        for i in li:
            i=i.replace('/','\\')
            li2.append(i)
        return li2
    else:
        return li


def replace_text(path,old_text, new_text):
    doc=Document(path)
    _,file_name=os.path.split(path)
    for p in doc.paragraphs:

            # 去除疼痛文字
            if p.style.font.size==95250:
                p.text=''
                print('{}:已去除不需要的信息'.format(file_name))

            inline = p.runs
            for i in inline:
                #run: element', 'font', 'italic', 'part', 'style', 'text', 'underline'
                #run.style 'base_style', 'builtin', 'delete', 'element', 'font', 'hidden', 'locked', 'name', 'part', 'priority', 'quick_style', 'style_id', 'type', 'unhide_when_used'
                # run.style.font  'all_caps', 'bold', 'color', 'complex_script', 'cs_bold', 'cs_italic', 'double_strike', 'element', 'emboss', 'hidden', 'highlight_color', 'imprint', 'italic', 'math', 'name', 'no_proof', 'outline', 'part', 'rtl', 'shadow', 'size', 'small_caps', 'snap_to_grid', 'spec_vanish', 'strike', 'subscript', 'superscript', 'underline', 'web_hidden'


                if old_text in i.text:
                    i.text = i.text.replace(old_text, new_text)

                    print('{}:已去除不需要的信息'.format(file_name))



def replace_text_page(path):
    doc=Document(path)
    _,file_name=os.path.split(path)
    for p in doc.paragraphs:
        old_text=re.findall('第\d+页|共\d+页',p.text)
        if old_text!=[]:
            for a in old_text:
                p.text = p.text.replace(a, '')
                print('{}:已去除不需要的信息'.format(file_name))
        text_2='  '
        p.text=replace_space(p.text,file_name,space_type=text_2)
        doc.save(path)

def replace_space(content,file_name,space_type=' '):
    if space_type in content:
        content=content.replace(space_type,'')
        print('{}:已去除不需要的信息'.format(file_name))
        replace_space(content,file_name,space_type)
    else:
        return content

def replace_many(path,filter_li,replace_to=''):
    for root,dirs,files in os.walk(path):
        for f in files:
            full_path=root+'/'+f
            replace_text_page(full_path)
            for replace_from in filter_li:
                replace_text(full_path,replace_from,replace_to)




def doc_to_docx(save_path,del_origin_file=False):
    for root,dirs,files in os.walk(save_path):
        word = wc.Dispatch("Word.Application") # 打开word应用程序
        for file in files:
            file_name,_=os.path.splitext(file)
            if file.endswith(('.doc','.DOC','.dotx')):
                full_path=root+'/'+file
                doc = word.Documents.Open(full_path) #打开word文件
                new_file_path="{}.docx".format(root+'/'+file_name)
                doc.SaveAs(new_file_path, 12)#另存为后缀为".docx"的文件，其中参数12指docx文件
                doc.Close() #关闭原来word文件
                if del_origin_file==True:
                    os.remove(full_path)
                print('{}:转换docx成功'.format(file_name))
        word.Quit()
        print("完成！")


def doc_to_pdf(src):
    if not os.path.exists(src):
        print(src + "不存在，无法继续！")
        return False
    # os.system('taskkill /im wps.exe')
    # 如果文件存在就删除
    dst=src.replace('.docx','.pdf')
    if os.path.exists(dst):
        os.remove(dst)
    o = wc.Dispatch("Kwps.Application")
    o.Visible = False
    doc = o.Documents.Open(src)
    doc.ExportAsFixedFormat(dst, 17)

    file_name=src.replace('.docx','')
    if os.path.exists(dst):
        print('{}:转换pdf完成'.format(file_name))

    else:
        print('{}:转换图片出错'.format(file_name))
    doc.Close()
    o.Quit()


def  many_doc_to_pdf(save_path):
    try:
        o = wc.Dispatch("Kwps.Application")
        o.Visible = False
        for root,dirs,files in os.walk(save_path):
            for file in files:
                changed=get_changed(save_path,choice=('.pdf',))
                pdf_file=os.path.splitext(file)[0]+'.pdf'
                if not pdf_file in changed:
                    full_path=root+'/'+file
                    if file.endswith('.docx'):
                        doc_to_pdf(full_path)
    except Exception as e:
        print(e)


def pdf_to_png(filename, outputDir):

        doc = fitz.open(filename)

        _,f_name=os.path.split(filename)
        f_name=f_name.replace('.pdf','')

        count=doc.pageCount

        if doc.pageCount>=5:
            count=5


        for pg in range(count):
            page = doc[pg]
            rotate = int(0)
            # 每个尺寸的缩放系数为2，这将为我们生成分辨率提高四倍的图像。
            zoom_x = 2.0
            zoom_y = 2.0
            trans = fitz.Matrix(zoom_x, zoom_y).preRotate(rotate)
            pm = page.getPixmap(matrix=trans, alpha=False)
            pm.writePNG('{}.png'.format(outputDir+'/'+str(pg)))
        print('{}:图片转换成功'.format(f_name))


def many_pdf_to_png(pdf_path,dir_path):
    for root,dirs,files in os.walk(pdf_path):
        for file in files:
            if file.endswith('.pdf'):
                full_path=root+'/'+file
                dir_name=file.replace('.pdf','')
                one_path=dir_path+'/'+dir_name
                if not os.path.exists(one_path):
                    os.mkdir(one_path)
                pdf_to_png(full_path,one_path)
                os.remove(full_path)



# def pdf_to_docx(pdf_path):
#     pdf_path,pdf_name=os.path.split(pdf_path)
#     # p_name,ext=os.path.splitext(pdf_name)
#     docx_path=pdf_path.replace('.pdf','docx')
#     content=pdf_to_str(pdf_path)
#     write_to_docx(content,docx_path)
#     print('{}:pdf转docx成功！'.format(pdf_name))


def rename_file(path,filter_word):
    for i in filter_word:
        for root,dirs,files in os.walk(path):
            for f in files:
                full_path=root+'/'+f
                if '+' in f:
                    new_name=f.replace('+','、')
                    new_path=root+'/'+new_name
                    os.rename(full_path,new_path)
                    print('{}：改名成功'.format(new_name))
                if ' ' in f:
                    new_name=f.replace(' ','')
                    new_path=root+'/'+new_name
                    if not os.path.exists(new_path):
                        os.rename(full_path,new_path)
                    else:
                        new_path,ext=os.path.splitext(new_path)
                        new_path=new_path+str(random.randint(1,1000))+ext
                        os.rename(full_path,new_path)
                    print('{}：改名成功'.format(new_name))

                if i in f:
                    new_name=f.replace(i,'')
                    new_path=root+'/'+new_name
                    os.rename(full_path,new_path)
                    print('{}：改名成功'.format(new_name))

                # filter_front(root,f)
                # filter_end(root,f)



def rename_file_front(path):

    for root,dirs,files in os.walk(path):
        for f in files:
            filter_front(root,f)

def rename_file_end(path):
    for root,dirs,files in os.walk(path):
        for f in files:
            filter_end(root,f)


def filter_front(root,f):
    full_path=root+'/'+f
    # \.|\-|、|\.|
    res=re.findall('^\d+\-\d+\-\d+|^\d+\-\d+|^\d+\.\d+\.\d+|^\d+\.\d+|^\d+\.|^\d+\、|^\d+\-|^\d+|^\d+[年|月]|^\d+',f)
    f_name,ext=os.path.splitext(f)

    if res!=[]:
        res=res[0]
        try:
            new_name=f_name.replace(res,'')
            new_path=root+'/'+new_name+ext
            os.rename(full_path,new_path)
        except Exception as e:
            print(e)
            filter_front(root,f)

        print('{}:改名成功'.format(f))

def filter_end(root,f):
    full_path=root+'/'+f
    f_name,ext=os.path.splitext(f)
    res2=re.findall('\(\d+\)|（\d+）|\d+\-\d+\-\d+$|\d+\-\d+$|\d+\.\d+\.\d$|\d+\.\d&|\d+\.$|\d+\、$|\d+\-$|\d+$|\d+[年|月]$|\d+$',f_name)
    if res2!=[]:
        res2=res2[0]
        try:
            new_name=f_name.replace(res2,random.choice(cfg.filter_after_word))
            new_path=root+'/'+new_name+ext
            os.rename(full_path,new_path)
        except Exception as e:
            print(e)
            filter_end(root,f)
        print('{}:改名成功'.format(new_name))

def rename_file_double(path,add_name=None):
    for root,dirs,files in os.walk(path):
        for f in files:
            full_name=os.path.join(root,f)
            dst_name,ext=os.path.splitext(full_name)
            if add_name!=None:
                dst_name=dst_name+str(add_name)
            else:
                dst_name=dst_name+str(random.randint(1,1000))
            final_name=dst_name+ext
            os.rename(full_name,final_name)
            print('{}:重命名成功'.format(dst_name))


if __name__ == "__main__":
    # path='f:/资源/WORD资源/2019/10月/待处理'
    # path='f:/资源/PPT资源/2019/10月/2019-10-18'
    # path='f:/资源/WORD资源/2019/10月/2019-10-18'
    # path='f:/资源/WORD资源/2019/10月/2019-10-18'
    # path='f:/资源/WORD资源/2019/10月/暂存1'
    rename_file_double(com.path)

    # replace_many(path,cfg.word_filter)
