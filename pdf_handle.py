#!/usr/bin/env python
# -*- coding: utf-8 -*-

__mtime__ = '2019/10/12'

import os
from configparser import ConfigParser
from io import StringIO
from io import open
from concurrent.futures import ProcessPoolExecutor

from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import process_pdf
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

import config as cfg


test_path='f:/资源/WORD资源/2019/10月/2019-10-12/2014年劳动合同范本.pdf'
to_path='f:/资源/WORD资源/2019/10月/2019-10-12/2014年劳动合同范本.docx'



def read_from_pdf(file_path):
    with open(file_path, 'rb') as file:
        resource_manager = PDFResourceManager()
        return_str = StringIO()
        lap_params = LAParams()

        device = TextConverter(
            resource_manager, return_str, laparams=lap_params)
        process_pdf(resource_manager, device, file)
        device.close()

        content = return_str.getvalue()
        return_str.close()
        return content


def save_text_to_word(content, file_path):
    doc = Document()
    # for line in content.split('\n')[1]:
    #     p = doc.add_paragraph()
    #     style = doc.styles.add_style('UserStyle%d' %line,WD_STYLE_TYPE.PARAGRAPH)
    #     style.font.size = Pt(32)
    #     p.style = style

    for line in content.split('\n')[:]:
        paragraph = doc.add_paragraph()
        doc.styles['Normal'].font.name = u'宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        paragraph.add_run(remove_control_characters(line))
    doc.save(file_path)


def remove_control_characters(content):
    mpa = dict.fromkeys(range(32))
    return content.translate(mpa)


def pdf_to_word(pdf_file_path, word_file_path):
    content = read_from_pdf(pdf_file_path)
    save_text_to_word(content, word_file_path)

def main():
    config_parser = ConfigParser()
    config_parser.read('config.cfg')
    # config = config_parser['default']

    tasks = []
    with ProcessPoolExecutor(max_workers=int(cfg.max_worker)) as executor:
        for file in os.listdir(cfg.save_path):
            extension_name = os.path.splitext(file)[1]
            if extension_name != '.pdf':
                continue
            file_name = os.path.splitext(file)[0]
            pdf_file =cfg.save_path + '/' + file
            word_file = cfg.save_path + '/' + file_name + '.docx'
            print('正在处理: ', file)
            result = executor.submit(pdf_to_word, pdf_file, word_file)
            tasks.append(result)
    while True:
        exit_flag = True
        for task in tasks:
            if not task.done():
                exit_flag = False
        if exit_flag:
            print('完成')
            exit(0)



if __name__=='__main__':
    main()

